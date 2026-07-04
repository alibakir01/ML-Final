"""Rebuild the mangled Table II (XGBoost hyper-parameters) cleanly."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

F = 'Final_Project_Paper_Bakir_Baskal_IEEE_v2.docx'
d = Document(F)

data = [
    ('Hyper parameter', 'Value'),
    ('n_estimators (cap)', '3000'),
    ('learning_rate', '0.0271'),
    ('max_depth', '6'),
    ('min_child_weight', '2'),
    ('subsample', '0.793'),
    ('colsample_bytree', '0.718'),
    ('gamma', '0.268'),
    ('reg_alpha', '0.0151'),
    ('reg_lambda', '3.401'),
    ('early_stopping_rounds', '50 (≈460 trees kept)'),
]

old = d.tables[0]
old_tbl = old._tbl
try:
    old_style = old.style
except Exception:
    old_style = None

# build replacement at end, then relocate
new = d.add_table(rows=len(data), cols=2)
new.style = old_style if old_style is not None else d.styles['Table Grid']
new.autofit = True

for i, (a, b) in enumerate(data):
    c0, c1 = new.cell(i, 0), new.cell(i, 1)
    c0.text = a
    c1.text = b
    for cell, align in ((c0, WD_ALIGN_PARAGRAPH.LEFT), (c1, WD_ALIGN_PARAGRAPH.CENTER)):
        for p in cell.paragraphs:
            p.alignment = align
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

# move new table to the old table's position, then delete old
old_tbl.addprevious(new._tbl)
old_tbl.getparent().remove(old_tbl)

d.save(F)
print('Table II rebuilt; tables now:', len(d.tables))
