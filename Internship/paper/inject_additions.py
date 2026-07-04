"""
inject_additions.py
Insert the peer-review additions (math critique, methodology equations,
significance test, residual analysis, DOIs) into the user's ORIGINAL document
(orig_editable.docx, obtained by Word's PDF->docx conversion), preserving its
format and content. Equations are produced as native Word OMML via pandoc and
spliced in; figures are inserted with python-docx; DOIs are appended to the
reference list.
"""
import copy
import pypandoc
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = 'orig_editable.docx'
DST = 'Final_Project_Paper_Bakir_Baskal_IEEE_v2.docx'
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

# ---- 1. block .tex -> .docx (native equations) --------------------------
blocks = {}
for name in ['A', 'B', 'C', 'D', 'E', 'F']:
    out = f'_blocks/block{name}.docx'
    pypandoc.convert_file(f'_blocks/block{name}.tex', 'docx', outputfile=out)
    blocks[name] = out
print('pandoc blocks built')

d = Document(SRC)


# ---- helpers ------------------------------------------------------------
def find_anchor(substr):
    for p in d.paragraphs:
        if substr in p.text:
            return p
    raise RuntimeError('anchor not found: ' + substr)


def is_math_only(p_elem):
    has_math = p_elem.find('.//{%s}oMath' % M_NS) is not None
    txt = ''.join(t.text or '' for t in p_elem.iter(qn('w:t')))
    return has_math and not txt.strip()


def style_inserted(para):
    """Apply Normal style + 11pt Times to text; centre pure-equation paras."""
    try:
        para.style = d.styles['Normal']
    except Exception:
        pass
    for r in para.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    if is_math_only(para._p):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def insert_block_after(last_elem, block_path):
    """Splice all non-empty paragraphs of block_path after last_elem."""
    bdoc = Document(block_path)
    for bp in bdoc.paragraphs:
        if not bp.text.strip() and not is_math_only(bp._p):
            continue
        cp = copy.deepcopy(bp._p)
        last_elem.addnext(cp)
        last_elem = cp
        style_inserted(Paragraph(cp, d._body))
    return last_elem


def insert_heading_after(last_elem, text):
    new_p = last_elem.makeelement(qn('w:p'), {})
    last_elem.addnext(new_p)
    para = Paragraph(new_p, d._body)
    try:
        para.style = d.styles['Heading 2']
    except Exception:
        pass
    para.add_run(text)
    return new_p


def insert_picture_after(last_elem, img, width_in=3.3):
    new_p = last_elem.makeelement(qn('w:p'), {})
    last_elem.addnext(new_p)
    para = Paragraph(new_p, d._body)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img, width=Inches(width_in))
    return new_p


def insert_caption_after(last_elem, text):
    new_p = last_elem.makeelement(qn('w:p'), {})
    last_elem.addnext(new_p)
    para = Paragraph(new_p, d._body)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(text)
    r.font.size = Pt(9)
    r.italic = True
    return new_p


def set_para_text(para, newtext):
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    para.add_run(newtext)


# ---- 2. Literature Review: linear-model math critique -------------------
a = find_anchor('are covered in Modules 10 and 12.')
last = insert_heading_after(a._p, 'E. Why Linear Models Fail on a Zero-Inflated, Bounded Target')
last = insert_block_after(last, blocks['A'])

# ---- 3. Methodology: target-encoder equations ---------------------------
a = find_anchor('so that no row ever sees its own target')
insert_block_after(a._p, blocks['B'])

# ---- 4. Methodology: MLP equations --------------------------------------
a = find_anchor('led us to expect the trees to win')
insert_block_after(a._p, blocks['C'])

# ---- 5. Methodology: SLSQP constrained objective ------------------------
a = find_anchor('while Gradient Boosting was dropped to zero')
insert_block_after(a._p, blocks['D'])

# ---- 6. Results: significance + residual subsections --------------------
a = find_anchor('shows the per model and blended OOF scores side by side')
last = insert_heading_after(a._p, 'C. Statistical Significance of the Blend Gain')
last = insert_block_after(last, blocks['E'])
last = insert_heading_after(last, 'D. Residual Analysis on the [0, 92] Boundaries')
last = insert_block_after(last, blocks['F'])
last = insert_picture_after(last, 'figures/fig6_residuals.png')
last = insert_caption_after(last,
    'Fig. 6. Residual diagnostics for the blend on the development OOF '
    'predictions: residual vs. fitted, residual vs. true target with the '
    'boundary bands highlighted, residual histogram, and normal Q-Q plot.')
last = insert_picture_after(last, 'figures/fig7_residual_bins.png')
last = insert_caption_after(last,
    'Fig. 7. Mean signed residual (y - y_hat) by true booked-nights bin. The '
    'blend over-predicts the low/zero range and under-predicts the high range.')

# renumber the existing Leaderboard subsection C -> E
for p in d.paragraphs:
    if p.text.strip() == 'C. Leaderboard result':
        set_para_text(p, 'E. Leaderboard result')
        p.style = d.styles['Heading 2']
        break

# ---- 7. References: append DOIs -----------------------------------------
def append_doi(substr, doi):
    for p in d.paragraphs:
        if substr in p.text:
            t = p.text.rstrip()
            t = (t[:-1] if t.endswith('.') else t) + f', doi: {doi}.'
            set_para_text(p, t)
            return True
    print('  DOI anchor not found:', substr)
    return False

# Wang reference is merged with Tang in one paragraph -> mid-text insert
for p in d.paragraphs:
    if 'Price determinants of sharing economy' in p.text and 'doi: 10.1016/j.ijhm' not in p.text:
        set_para_text(p, p.text.replace('pp. 120–131, 2017.',
                       'pp. 120–131, 2017, doi: 10.1016/j.ijhm.2016.12.007.', 1))
        break

doi_map = [
    ('Mapping Airbnb supply in European cities', '10.1016/j.annals.2018.02.008'),
    ('Determinants of Airbnb demand in Vienna', '10.1177/1354816617731196'),
    ('Computational Management Science, vol. 21', '10.1007/s10287-024-00513-9'),
    ('Applications, vol. 7, p. 100208, 2022', '10.1016/j.mlwa.2021.100208'),
    ('Induction of decision trees', '10.1007/BF00116251'),
    ('Bagging predictors', '10.1007/BF00058655'),
    ('Random forests', '10.1023/A:1010933404324'),
    ('Greedy function approximation', '10.1214/aos/1013203451'),
    ('XGBoost: A scalable tree boosting system', '10.1145/2939672.2939785'),
    ('The Elements of Statistical Learning', '10.1007/978-0-387-84858-7'),
    ('Stacked generalization', '10.1016/S0893-6080(05)80023-1'),
]
for substr, doi in doi_map:
    append_doi(substr, doi)

d.save(DST)
print('saved', DST)
