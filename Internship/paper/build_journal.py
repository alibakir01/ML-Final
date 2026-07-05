"""
build_ieee_docx.py
Convert the pandoc base docx (_journal_base.docx) into a true IEEE two-column
Word document, then it is exported to PDF by Word separately.

Layout produced:
  * US-Letter, IEEE margins (0.625in sides), single-column title block,
    two-column body via a continuous section break;
  * Times New Roman 10pt justified body;
  * IEEE heading numbering  I. / A. / 1)  with centred small-caps level-1,
    italic level-2 headings;
  * figures capped to column width; captions at 8pt.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = '_journal_base.docx'
DST = 'Final_Project_Paper_Bakir_Baskal_JOURNAL.docx'
TW = 1440  # twips per inch

d = Document(SRC)

# ---------------------------------------------------------------- font: Normal
def set_font(style, name='Times New Roman', size=10):
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), name)

for sn in ('Normal', 'Body Text', 'First Paragraph'):
    try:
        set_font(d.styles[sn])
    except KeyError:
        pass
d.styles['Normal'].paragraph_format.space_after = Pt(0)
d.styles['Normal'].paragraph_format.line_spacing = 1.0

# ---------------------------------------------------------------- page geometry
sec = d.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(0.625); sec.right_margin = Inches(0.625)

# ---------------------------------------------------------------- helpers
def delete_para(p):
    p._element.getparent().remove(p._element)

def make_cols(num, space=360, equal=True):
    c = OxmlElement('w:cols')
    c.set(qn('w:num'), str(num)); c.set(qn('w:space'), str(space))
    if equal:
        c.set(qn('w:equalWidth'), '1')
    return c

def build_sectPr(cols_num, continuous=True):
    sp = OxmlElement('w:sectPr')
    if continuous:
        t = OxmlElement('w:type'); t.set(qn('w:val'), 'continuous'); sp.append(t)
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(int(8.5 * TW))); pgSz.set(qn('w:h'), str(int(11 * TW)))
    sp.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    for k, v in [('top', 0.75), ('right', 0.625), ('bottom', 1.0), ('left', 0.625),
                 ('header', 0.5), ('footer', 0.5), ('gutter', 0)]:
        pgMar.set(qn('w:' + k), str(int(v * TW)))
    sp.append(pgMar)
    sp.append(make_cols(cols_num))
    return sp

# ---------------------------------------------------------------- title block
title_p = d.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title_p.runs:
    r.font.size = Pt(20); r.font.bold = True; r.font.name = 'Times New Roman'

# delete the stray \markboth running-head paragraph
for p in list(d.paragraphs[:8]):
    if p.text.strip().startswith('Bak') and 'Predicting Quarterly Airbnb Booking Counts in NYC' in p.text:
        delete_para(p)

paras = d.paragraphs

# author
for p in paras[:6]:
    if p.style.name == 'Author':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(11); r.font.name = 'Times New Roman'

# Abstract title + body; keywords -> Index Terms
def first_with_style(name):
    for p in d.paragraphs:
        if p.style.name == name:
            return p
    return None

ab_title = first_with_style('Abstract Title')
if ab_title:
    for r in ab_title.runs:
        r.font.bold = True; r.font.size = Pt(9)
ab_body = first_with_style('Abstract')
if ab_body:
    for r in ab_body.runs:
        r.font.italic = True; r.font.size = Pt(9); r.font.bold = True

# keywords paragraph = the Body Text just before the first Heading 1
kw_para = None
for i, p in enumerate(d.paragraphs):
    if p.style.name == 'Heading 1':
        # walk back to previous non-empty paragraph
        j = i - 1
        while j >= 0 and not d.paragraphs[j].text.strip():
            j -= 1
        kw_para = d.paragraphs[j]
        break
if kw_para is not None and not kw_para.text.strip().lower().startswith('index terms'):
    for r in kw_para.runs:
        r.font.italic = True; r.font.size = Pt(9)
    run = kw_para.runs[0]
    run.text = 'Index Terms—' + run.text
    run.font.bold = True

# ---------------------------------------------------------------- section break
# title-block section (1 col) ends at kw_para; body becomes 2 cols.
if kw_para is not None:
    pPr = kw_para._p.get_or_add_pPr()
    pPr.append(build_sectPr(1, continuous=True))

# final body sectPr -> 2 columns, CONTINUOUS so the body flows on page 1
body_sp = d.sections[-1]._sectPr
# (a) continuous type, placed in schema-correct position (immediately before pgSz,
#     i.e. after any header/footer/footnote references)
for t in body_sp.findall(qn('w:type')):
    body_sp.remove(t)
btype = OxmlElement('w:type'); btype.set(qn('w:val'), 'continuous')
pgSz_el = body_sp.find(qn('w:pgSz'))
if pgSz_el is not None:
    pgSz_el.addprevious(btype)
else:
    body_sp.insert(0, btype)
# (b) replace cols with a 2-column spec, placed before w:docGrid
for c in body_sp.findall(qn('w:cols')):
    body_sp.remove(c)
cols2 = make_cols(2)
docGrid = body_sp.find(qn('w:docGrid'))
if docGrid is not None:
    docGrid.addprevious(cols2)
else:
    body_sp.append(cols2)

# ---------------------------------------------------------------- headings + numbering
roman = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']
letters = [chr(c) for c in range(ord('A'), ord('Z') + 1)]
h1 = h2 = h3 = 0

def style_runs(p, size=10, bold=False, italic=False, smallcaps=False):
    for r in p.runs:
        r.font.size = Pt(size); r.font.name = 'Times New Roman'
        r.font.bold = bold; r.font.italic = italic
        r.font.small_caps = smallcaps

def prepend(p, text):
    if p.runs:
        p.runs[0].text = text + p.runs[0].text
    else:
        p.add_run(text)

for p in d.paragraphs:
    sn = p.style.name
    if sn == 'Heading 1':
        h2 = 0; h3 = 0
        prepend(p, f'{roman[h1]}. ')
        h1 += 1
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_runs(p, size=10, bold=False, smallcaps=True)
    elif sn == 'Heading 2':
        h3 = 0
        prepend(p, f'{letters[h2]}. ')
        h2 += 1
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_runs(p, size=10, italic=True)
    elif sn == 'Heading 3':
        prepend(p, f'{h3 + 1}) ')
        h3 += 1
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_runs(p, size=10, italic=True)
    elif sn in ('Heading 4', 'Heading 5', 'Heading 6'):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_runs(p, size=10, italic=True, bold=True)

# ---------------------------------------------------------------- captions 8pt
for p in d.paragraphs:
    if 'Caption' in p.style.name or p.style.name in ('Image Caption', 'Table Caption'):
        for r in p.runs:
            r.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------------------------------------------------------- justify body
for p in d.paragraphs:
    if p.style.name in ('Body Text', 'First Paragraph', 'Normal'):
        if p.alignment is None:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ---------------------------------------------------------------- shrink figures to column width
COLW = Emu(int(3.30 * 914400))  # ~3.30 in column
for shp in d.inline_shapes:
    w, h = shp.width, shp.height
    if w and w > COLW:
        ratio = COLW / w
        shp.width = int(w * ratio)
        shp.height = int(h * ratio)

# ---------------------------------------------------------------- in-text "Section N" -> Roman
sec_map = {str(i + 1): roman[i] for i in range(7)}
pat = re.compile(r'Section(\s| )([1-7])\b')
for p in d.paragraphs:
    for r in p.runs:
        if 'Section' in r.text:
            r.text = pat.sub(lambda m: 'Section ' + sec_map[m.group(2)], r.text)

# ---------------------------------------------------------------- feature-inventory table widths
# The 2-column "Complete Feature Inventory" longtable flows in the normal
# 2-column body (~3.30in per column), but pandoc doesn't reliably carry over
# the LaTeX p{2.2in}p{0.8in} column spec, so set it explicitly here with a
# FIXED layout (otherwise Word autofits and wraps "float"/"integer" badly).
inv_table = d.tables[-1]
tbl = inv_table._tbl
tblPr = tbl.tblPr
for old in tblPr.findall(qn('w:tblW')):
    tblPr.remove(old)
tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'), str(int(3.15 * TW)))
tblW.set(qn('w:type'), 'dxa')
tblPr.append(tblW)
for old in tblPr.findall(qn('w:tblLayout')):
    tblPr.remove(old)
tblLayout = OxmlElement('w:tblLayout')
tblLayout.set(qn('w:type'), 'fixed')
tblPr.append(tblLayout)
NAME_W = Emu(int(2.35 * 914400))
TYPE_W = Emu(int(0.80 * 914400))
for col, w in zip(inv_table.columns, [NAME_W, TYPE_W]):
    col.width = w
for row in inv_table.rows:
    for cell, w in zip(row.cells, [NAME_W, TYPE_W]):
        cell.width = w

d.save(DST)
print('Saved', DST)
